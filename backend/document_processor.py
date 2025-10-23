"""
Document Processing Module
Handles extraction of text from PDF and DOCX files for the interview system.
"""

import logging
from typing import Optional
from pathlib import Path
import PyPDF2
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Custom exception for document processing errors."""
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content

    Raises:
        DocumentProcessingError: If PDF extraction fails
    """
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            logger.info(f"Processing PDF with {num_pages} pages")

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            if not text.strip():
                raise DocumentProcessingError("PDF file appears to be empty or text extraction failed")

            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()

    except FileNotFoundError:
        error_msg = f"PDF file not found: {file_path}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg)
    except PyPDF2.errors.PdfReadError as e:
        error_msg = f"Invalid or corrupted PDF file: {str(e)}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg)
    except Exception as e:
        error_msg = f"Unexpected error processing PDF: {str(e)}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content

    Raises:
        DocumentProcessingError: If DOCX extraction fails
    """
    try:
        doc = Document(file_path)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"

        if not text.strip():
            raise DocumentProcessingError("DOCX file appears to be empty")

        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text.strip()

    except FileNotFoundError:
        error_msg = f"DOCX file not found: {file_path}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg)
    except Exception as e:
        error_msg = f"Unexpected error processing DOCX: {str(e)}"
        logger.error(error_msg)
        raise DocumentProcessingError(error_msg)


def extract_text_from_document(file_path: str, file_type: Optional[str] = None) -> str:
    """
    Extract text from a document (PDF or DOCX).
    Automatically detects file type if not provided.

    Args:
        file_path: Path to the document file
        file_type: Optional file type ('pdf' or 'docx'). If None, will detect from extension.

    Returns:
        Extracted text content

    Raises:
        DocumentProcessingError: If extraction fails or file type is unsupported
    """
    if file_type is None:
        file_extension = Path(file_path).suffix.lower()
        if file_extension == '.pdf':
            file_type = 'pdf'
        elif file_extension in ['.docx', '.doc']:
            file_type = 'docx'
        else:
            raise DocumentProcessingError(f"Unsupported file extension: {file_extension}")

    file_type = file_type.lower()

    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    else:
        raise DocumentProcessingError(f"Unsupported file type: {file_type}. Only 'pdf' and 'docx' are supported.")
